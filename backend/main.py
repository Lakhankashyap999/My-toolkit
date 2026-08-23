import fitz  # PyMuPDF
from fastapi import FastAPI, UploadFile, File, Form
from fastapi.responses import Response
from docx import Document
from docx.shared import Pt
import io

app = FastAPI()

@app.post("/convert")
async def convert_pdf_to_word(file: UploadFile = File(...)):
    pdf_bytes = await file.read()
    pdf_doc = fitz.open(stream=pdf_bytes, filetype="pdf")

    docx_doc = Document()

    for page_num in range(len(pdf_doc)):
        page = pdf_doc[page_num]
        blocks = page.get_text("dict")["blocks"]

        for block in blocks:
            if block["type"] == 0:
                for line in block["lines"]:
                    line_text = ""
                    max_font_size = 0
                    for span in line["spans"]:
                        line_text += span["text"]
                        if span["size"] > max_font_size:
                            max_font_size = span["size"]
                    if line_text.strip():
                        if max_font_size >= 16:
                            docx_doc.add_heading(line_text.strip(), level=1)
                        else:
                            p = docx_doc.add_paragraph(line_text.strip())
                            p.style.font.size = Pt(max_font_size)

        try:
            tables = page.find_tables()
            for table in tables:
                table_data = table.extract()
                if table_data:
                    rows = len(table_data)
                    cols = len(table_data[0]) if rows > 0 else 0
                    docx_table = docx_doc.add_table(rows=rows, cols=cols)
                    for i, row in enumerate(table_data):
                        for j, cell_text in enumerate(row):
                            docx_table.cell(i, j).text = cell_text or ""
        except Exception as e:
            print(f"Table extraction error: {e}")

    output = io.BytesIO()
    docx_doc.save(output)
    output.seek(0)

    return Response(
        content=output.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=converted.docx"}
    )

@app.post("/edit")
async def edit_pdf(
    file: UploadFile = File(...),
    add_text: str = Form(""),
    add_text_x: float = Form(50),
    add_text_y: float = Form(50),
    text_size: int = Form(16),
    text_color: str = Form("#000000"),
    watermark: str = Form(""),
    watermark_opacity: float = Form(0.3),
    watermark_size: int = Form(40),
    delete_pages: str = Form(""),
    rotate_angle: int = Form(0),
    selected_page: int = Form(1),
):
    pdf_bytes = await file.read()
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")

    # Rotate all pages if needed
    if rotate_angle != 0:
        for page in doc:
            page.set_rotation((page.rotation + rotate_angle) % 360)

    # Delete pages (indices descending)
    if delete_pages.strip():
        pages_to_delete = []
        parts = delete_pages.split(",")
        for part in parts:
            part = part.strip()
            if "-" in part:
                start, end = part.split("-")
                start, end = int(start), int(end)
                pages_to_delete.extend(range(start, end + 1))
            else:
                pages_to_delete.append(int(part))
        # Convert to zero-based, deduplicate, sort descending
        indices = sorted(set(p - 1 for p in pages_to_delete), reverse=True)
        for idx in indices:
            if 0 <= idx < len(doc):
                doc.delete_page(idx)

    # Add watermark to all pages
    if watermark.strip():
        for page in doc:
            rect = page.rect
            font_size = watermark_size
            text_width = fitz.get_text_length(watermark, fontname="helv", fontsize=font_size)
            x = (rect.width - text_width) / 2
            y = rect.height / 2
            page.insert_text(
                (x, y),
                watermark,
                fontsize=font_size,
                fontname="helv",
                color=(0.5, 0.5, 0.5),
                alpha=watermark_opacity,
                rotate=45,
            )

    # Add text on selected page
    if add_text.strip():
        if 0 <= selected_page - 1 < len(doc):
            page = doc[selected_page - 1]
            # Convert color hex to tuple
            hex_color = text_color.lstrip("#")
            r = int(hex_color[0:2], 16) / 255
            g = int(hex_color[2:4], 16) / 255
            b = int(hex_color[4:6], 16) / 255
            page.insert_text(
                (add_text_x, add_text_y),
                add_text,
                fontsize=text_size,
                fontname="helv",
                color=(r, g, b),
            )

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)

    return Response(
        content=output.getvalue(),
        media_type="application/pdf",
        headers={"Content-Disposition": 'attachment; filename="edited.pdf"'}
    )